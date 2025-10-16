"""
Datenverarbeitungsmodul für Verwaltungstexte
Speziell optimiert für deutsche Verwaltungs- und Rechtsterminologie.
"""

import json
import pandas as pd
from pathlib import Path
from typing import List, Dict, Any, Optional
from datasets import Dataset, load_dataset
from transformers import PreTrainedTokenizer
import re
import logging

# PDF und Markdown Verarbeitung
try:
    import PyPDF2
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import markdown
    from bs4 import BeautifulSoup
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)

class VerwaltungsDataProcessor:
    """
    Datenverarbeitung für deutsche Verwaltungstexte.
    Unterstützt verschiedene Eingabeformate und bereitet sie für das Training vor.
    """
    
    def __init__(self, tokenizer: PreTrainedTokenizer, max_length: int = 512):
        self.tokenizer = tokenizer
        self.max_length = max_length
        self.verwaltungs_keywords = {
            'behörden', 'verwaltung', 'amt', 'bescheid', 'antrag', 'verordnung',
            'gesetz', 'paragraph', 'artikel', 'recht', 'rechtlich', 'zuständig',
            'bearbeitung', 'verfahren', 'frist', 'genehmigung', 'bewilligung',
            'steuer', 'abgabe', 'gebühr', 'bußgeld', 'ordnungswidrigkeit'
        }
    
    def load_and_process(self, data_path: str) -> Dataset:
        """
        Lade und verarbeite Daten aus verschiedenen Quellen.
        
        Args:
            data_path: Pfad zur Datendatei (txt, json, jsonl, csv) oder Verzeichnis mit vielen Dateien
            
        Returns:
            Dataset: Verarbeiteter Hugging Face Dataset
        """
        data_path = Path(data_path)
        
        if not data_path.exists():
            raise FileNotFoundError(f"Pfad nicht gefunden: {data_path}")
        
        # Prüfe ob es ein Verzeichnis ist
        if data_path.is_dir():
            texts = self._load_directory(data_path)
        else:
            # Einzelne Datei laden
            if data_path.suffix.lower() == '.txt':
                texts = self._load_txt(data_path)
            elif data_path.suffix.lower() == '.json':
                texts = self._load_json(data_path)
            elif data_path.suffix.lower() == '.jsonl':
                texts = self._load_jsonl(data_path)
            elif data_path.suffix.lower() == '.csv':
                texts = self._load_csv(data_path)
            elif data_path.suffix.lower() == '.pdf':
                texts = self._load_pdf(data_path)
            elif data_path.suffix.lower() in ['.md', '.markdown']:
                texts = self._load_markdown(data_path)
            elif data_path.suffix.lower() in ['.docx', '.doc']:
                texts = self._load_docx(data_path)
            else:
                raise ValueError(f"Unsupported file format: {data_path.suffix}")
        
        logger.info(f"Loaded {len(texts)} texts from {data_path}")
        
        # Texte vorverarbeiten
        processed_texts = self._preprocess_texts(texts)
        
        # In Dataset konvertieren
        dataset = Dataset.from_dict({"text": processed_texts})
        
        # Tokenisierung
        dataset = dataset.map(
            self._tokenize_function,
            batched=True,
            remove_columns=["text"]
        )
        
        return dataset
    
    def _load_txt(self, path: Path) -> List[str]:
        """Lade Texte aus einer einfachen Textdatei."""
        with open(path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Aufteilen in Absätze oder Sätze
        texts = [text.strip() for text in content.split('\n\n') if text.strip()]
        return texts
    
    def _load_json(self, path: Path) -> List[str]:
        """Lade Texte aus einer JSON-Datei."""
        with open(path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        if isinstance(data, list):
            if isinstance(data[0], str):
                return data
            elif isinstance(data[0], dict):
                return [item.get('text', '') for item in data]
        elif isinstance(data, dict):
            if 'texts' in data:
                return data['texts']
            elif 'data' in data:
                return data['data']
        
        raise ValueError("JSON format not recognized. Expected list of strings or dicts with 'text' field.")
    
    def _load_jsonl(self, path: Path) -> List[str]:
        """Lade Texte aus einer JSONL-Datei."""
        texts = []
        with open(path, 'r', encoding='utf-8') as f:
            for line in f:
                data = json.loads(line.strip())
                if isinstance(data, str):
                    texts.append(data)
                elif isinstance(data, dict):
                    texts.append(data.get('text', ''))
        return texts
    
    def _load_csv(self, path: Path) -> List[str]:
        """Lade Texte aus einer CSV-Datei."""
        df = pd.read_csv(path)
        
        # Versuche verschiedene Spaltennamen
        text_columns = ['text', 'content', 'document', 'body', 'message']
        
        for col in text_columns:
            if col in df.columns:
                return df[col].fillna('').tolist()
        
        # Falls keine passende Spalte gefunden, nimm die erste Textspalte
        string_cols = df.select_dtypes(include=['object']).columns
        if len(string_cols) > 0:
            return df[string_cols[0]].fillna('').tolist()
        
        raise ValueError("No text column found in CSV file.")
    
    def _load_pdf(self, path: Path) -> List[str]:
        """Lade Text aus PDF-Dateien."""
        if not PDF_AVAILABLE:
            logger.warning(f"PDF-Bibliotheken nicht verfügbar. Überspringe {path}")
            return []
        
        texts = []
        
        try:
            # Versuche zuerst pdfplumber (bessere OCR-Unterstützung)
            with pdfplumber.open(path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    text = page.extract_text()
                    if text and text.strip():
                        # Seitennummer als Metadaten hinzufügen
                        clean_text = self._clean_pdf_text(text)
                        if len(clean_text.strip()) > 50:  # Nur substanzielle Seiten
                            texts.append(f"[Seite {page_num + 1}]\n{clean_text}")
                        
        except Exception as e:
            logger.warning(f"pdfplumber fehlgeschlagen für {path}: {e}")
            
            # Fallback zu PyPDF2
            try:
                with open(path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page_num, page in enumerate(pdf_reader.pages):
                        text = page.extract_text()
                        if text and text.strip():
                            clean_text = self._clean_pdf_text(text)
                            if len(clean_text.strip()) > 50:
                                texts.append(f"[Seite {page_num + 1}]\n{clean_text}")
                                
            except Exception as e2:
                logger.error(f"Beide PDF-Reader fehlgeschlagen für {path}: {e2}")
                return []
        
        logger.info(f"Extrahiert {len(texts)} Seiten aus PDF: {path.name}")
        return texts
    
    def _clean_pdf_text(self, text: str) -> str:
        """Bereinige PDF-Text von typischen Artefakten."""
        # Mehrfache Leerzeichen/Newlines reduzieren
        text = re.sub(r'\n+', '\n', text)
        text = re.sub(r' +', ' ', text)
        
        # Typische PDF-Artefakte entfernen
        text = re.sub(r'\x0c', '', text)  # Form Feed
        text = re.sub(r'[^\w\s\-.,;:!?()\[\]§""„"„"\'°%&/=+*#]', ' ', text)  # Seltsame Zeichen
        
        # Zeilenumbrüche in Sätzen korrigieren
        text = re.sub(r'([a-z])\n([a-z])', r'\1 \2', text)
        
        return text.strip()
    
    def _load_markdown(self, path: Path) -> List[str]:
        """Lade und verarbeite Markdown-Dateien."""
        if not MARKDOWN_AVAILABLE:
            logger.warning(f"Markdown-Bibliotheken nicht verfügbar. Überspringe {path}")
            return []
        
        try:
            with open(path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Markdown zu HTML konvertieren
            html = markdown.markdown(content, extensions=['extra', 'codehilite'])
            
            # HTML zu Text konvertieren
            soup = BeautifulSoup(html, 'html.parser')
            
            texts = []
            
            # Verschiedene Abschnitte extrahieren
            # Headers als Kontext beibehalten
            current_section = ""
            
            for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol', 'blockquote']):
                if element.name.startswith('h'):
                    # Header - als neuer Kontext
                    current_section = element.get_text().strip()
                else:
                    # Content
                    text = element.get_text().strip()
                    if text and len(text) > 20:
                        if current_section:
                            full_text = f"[{current_section}]\n{text}"
                        else:
                            full_text = text
                        texts.append(full_text)
            
            # Falls keine strukturierten Elemente, nimm rohen Text
            if not texts:
                # Entferne Markdown-Syntax
                plain_text = re.sub(r'[#*_`\[\]()]', '', content)
                texts = [plain_text.strip()]
            
            logger.info(f"Extrahiert {len(texts)} Abschnitte aus Markdown: {path.name}")
            return texts
            
        except Exception as e:
            logger.error(f"Fehler beim Laden von Markdown {path}: {e}")
            return []
    
    def _load_docx(self, path: Path) -> List[str]:
        """Lade Text aus Word-Dokumenten (.docx)."""
        if not DOCX_AVAILABLE:
            logger.warning(f"python-docx nicht verfügbar. Überspringe {path}")
            return []
        
        try:
            doc = Document(path)
            texts = []
            
            # Absätze extrahieren
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text and len(text) > 10:
                    # Stil-Information hinzufügen falls verfügbar
                    style_name = paragraph.style.name if paragraph.style else ""
                    if style_name and "heading" in style_name.lower():
                        text = f"[{style_name}] {text}"
                    texts.append(text)
            
            # Tabellen extrahieren
            for table in doc.tables:
                table_text = ""
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        table_text += row_text + "\n"
                
                if table_text.strip():
                    texts.append(f"[Tabelle]\n{table_text.strip()}")
            
            logger.info(f"Extrahiert {len(texts)} Elemente aus Word-Dokument: {path.name}")
            return texts
            
        except Exception as e:
            logger.error(f"Fehler beim Laden von Word-Dokument {path}: {e}")
            return []
    
    def _load_directory(self, directory_path: Path) -> List[str]:
        """
        Lade alle unterstützten Dateien aus einem Verzeichnis rekursiv.
        Optimiert für große Verzeichnisse mit vielen Dateien.
        
        Args:
            directory_path: Pfad zum Verzeichnis
            
        Returns:
            List[str]: Alle geladenen Texte
        """
        
        supported_extensions = {'.txt', '.json', '.jsonl', '.csv', '.md', '.markdown', '.pdf', '.docx', '.doc'}
        all_texts = []
        
        logger.info(f"Durchsuche Verzeichnis: {directory_path}")
        
        # Verfügbarkeit der verschiedenen Formate prüfen
        available_extensions = {'.txt', '.json', '.jsonl', '.csv'}
        if PDF_AVAILABLE:
            available_extensions.add('.pdf')
        if MARKDOWN_AVAILABLE:
            available_extensions.update({'.md', '.markdown'})
        if DOCX_AVAILABLE:
            available_extensions.update({'.docx', '.doc'})
        
        # Nur verfügbare Extensions verwenden
        extensions_to_process = supported_extensions.intersection(available_extensions)
        
        logger.info(f"Unterstützte Formate: {', '.join(extensions_to_process)}")
        
        # Alle Dateien rekursiv finden
        all_files = []
        for ext in extensions_to_process:
            # Glob-Pattern für rekursive Suche
            pattern = f"**/*{ext}"
            files = list(directory_path.glob(pattern))
            all_files.extend(files)
        
        logger.info(f"Gefundene Dateien: {len(all_files)}")
        
        # Fortschrittsanzeige für große Verzeichnisse
        processed_count = 0
        error_count = 0
        
        for file_path in all_files:
            try:
                # Sehr große Dateien überspringen (> 50MB)
                if file_path.stat().st_size > 50 * 1024 * 1024:
                    logger.warning(f"Überspringe große Datei: {file_path} ({file_path.stat().st_size / 1024 / 1024:.1f}MB)")
                    continue
                
                # Basierend auf Dateierweiterung laden
                ext = file_path.suffix.lower()
                if ext in ['.txt', '.md', '.markdown']:
                    if ext == '.txt':
                        texts = self._load_txt(file_path)
                    else:
                        texts = self._load_markdown(file_path)
                elif ext == '.json':
                    texts = self._load_json(file_path)
                elif ext == '.jsonl':
                    texts = self._load_jsonl(file_path)
                elif ext == '.csv':
                    texts = self._load_csv(file_path)
                elif ext == '.pdf':
                    texts = self._load_pdf(file_path)
                elif ext in ['.docx', '.doc']:
                    texts = self._load_docx(file_path)
                else:
                    continue
                
                # Gültige Texte hinzufügen
                valid_texts = [t for t in texts if t and len(t.strip()) > 10]
                all_texts.extend(valid_texts)
                
                processed_count += 1
                
                # Fortschritt alle 100 Dateien anzeigen
                if processed_count % 100 == 0:
                    logger.info(f"Verarbeitet: {processed_count}/{len(all_files)} Dateien, {len(all_texts)} Texte gesammelt")
                
            except Exception as e:
                error_count += 1
                logger.warning(f"Fehler beim Laden von {file_path}: {e}")
                
                # Bei zu vielen Fehlern abbrechen
                if error_count > 50:
                    logger.error("Zu viele Dateifehler, breche ab")
                    break
        
        logger.info(f"Verzeichnis-Verarbeitung abgeschlossen:")
        logger.info(f"- Verarbeitete Dateien: {processed_count}")
        logger.info(f"- Fehlerhafte Dateien: {error_count}")
        logger.info(f"- Gesammelte Texte: {len(all_texts)}")
        
        # Duplikate entfernen (optional)
        if len(all_texts) > 1000:  # Nur bei vielen Texten
            logger.info("Entferne Duplikate...")
            unique_texts = list(set(all_texts))
            logger.info(f"Duplikate entfernt: {len(all_texts)} -> {len(unique_texts)}")
            all_texts = unique_texts
        
        return all_texts
    
    def _preprocess_texts(self, texts: List[str]) -> List[str]:
        """
        Vorverarbeitung der Texte speziell für Verwaltungssprache.
        
        Args:
            texts: Liste von Rohtexten
            
        Returns:
            List[str]: Vorverarbeitete Texte
        """
        processed = []
        
        for text in texts:
            # Grundlegende Bereinigung
            text = self._clean_text(text)
            
            # Mindestlänge prüfen
            if len(text.split()) < 5:
                continue
            
            # Maximale Länge für Training
            if len(text) > self.max_length * 4:  # Grober Schätzwert für Token
                # Text in Chunks aufteilen
                chunks = self._split_into_chunks(text)
                processed.extend(chunks)
            else:
                processed.append(text)
        
        logger.info(f"Preprocessed {len(processed)} text chunks")
        return processed
    
    def _clean_text(self, text: str) -> str:
        """Grundlegende Textbereinigung."""
        # Mehrfache Leerzeichen/Newlines entfernen
        text = re.sub(r'\s+', ' ', text)
        
        # Spezielle Zeichen normalisieren
        text = text.replace('§', 'Paragraph ')
        text = text.replace('Art.', 'Artikel')
        text = text.replace('bzw.', 'beziehungsweise')
        text = text.replace('z.B.', 'zum Beispiel')
        text = text.replace('u.a.', 'unter anderem')
        text = text.replace('etc.', 'et cetera')
        
        # Führende/nachfolgende Leerzeichen entfernen
        text = text.strip()
        
        return text
    
    def _split_into_chunks(self, text: str, overlap: int = 50) -> List[str]:
        """Teile langen Text in überlappende Chunks auf."""
        words = text.split()
        chunk_size = self.max_length // 2  # Konservative Schätzung
        
        chunks = []
        for i in range(0, len(words), chunk_size - overlap):
            chunk = ' '.join(words[i:i + chunk_size])
            if len(chunk.split()) >= 5:  # Mindestlänge
                chunks.append(chunk)
        
        return chunks
    
    def _tokenize_function(self, examples):
        """Tokenisierungsfunktion für das Dataset."""
        # Text tokenisieren
        tokenized = self.tokenizer(
            examples["text"],
            truncation=True,
            padding=False,
            max_length=self.max_length,
            return_special_tokens_mask=True,
        )
        
        # Labels für Causal LM setzen (gleich input_ids)
        tokenized["labels"] = tokenized["input_ids"].copy()
        
        return tokenized
    
    def create_sample_data(self, output_path: str, num_samples: int = 100):
        """
        Erstelle Beispieldaten für Verwaltungstexte.
        
        Args:
            output_path: Pfad für die Ausgabedatei
            num_samples: Anzahl der zu erstellenden Beispiele
        """
        
        verwaltungs_templates = [
            "Das Amt für {bereich} teilt mit, dass Ihr Antrag vom {datum} bearbeitet wird. Die Bearbeitungszeit beträgt {frist} Werktage.",
            "Gemäß § {paragraph} der {verordnung} ist eine {genehmigung} erforderlich. Bitte reichen Sie die vollständigen Unterlagen ein.",
            "Ihr Bescheid vom {datum} bezüglich {thema} wurde erstellt. Die Gebühr beträgt {betrag} Euro und ist bis zum {frist} zu entrichten.",
            "Die Verwaltung weist darauf hin, dass {vorschrift} zu beachten ist. Bei Nichteinhaltung droht ein Bußgeld von bis zu {betrag} Euro.",
            "Für die Beantragung einer {genehmigung} sind folgende Unterlagen erforderlich: {unterlagen}. Die Bearbeitungsgebühr beträgt {betrag} Euro."
        ]
        
        bereiche = ["Bauwesen", "Gewerbeaufsicht", "Umweltschutz", "Verkehrswesen", "Soziales"]
        verordnungen = ["Bauordnung", "Gewerbeordnung", "Umweltschutzverordnung", "Straßenverkehrsordnung"]
        genehmigungen = ["Baugenehmigung", "Gewerbeanmeldung", "Umweltgenehmigung", "Sondergenehmigung"]
        
        import random
        
        sample_texts = []
        for _ in range(num_samples):
            template = random.choice(verwaltungs_templates)
            
            text = template.format(
                bereich=random.choice(bereiche),
                datum=f"{random.randint(1, 28):02d}.{random.randint(1, 12):02d}.2024",
                frist=random.randint(5, 30),
                paragraph=random.randint(1, 100),
                verordnung=random.choice(verordnungen),
                genehmigung=random.choice(genehmigungen),
                thema=random.choice(["Bauvorhaben", "Gewerbetätigkeit", "Umweltschutzmaßnahme"]),
                betrag=random.randint(50, 500),
                unterlagen="Personalausweis, Nachweis der Qualifikation, Versicherungsnachweis"
            )
            sample_texts.append(text)
        
        # Als JSONL speichern
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            for text in sample_texts:
                json.dump({"text": text}, f, ensure_ascii=False)
                f.write('\n')
        
        logger.info(f"Created {num_samples} sample texts in {output_path}")

if __name__ == "__main__":
    # Beispiel für die Verwendung
    from transformers import AutoTokenizer
    
    tokenizer = AutoTokenizer.from_pretrained("microsoft/DialoGPT-medium")
    processor = VerwaltungsDataProcessor(tokenizer)
    
    # Beispieldaten erstellen
    processor.create_sample_data("data/examples/verwaltung_sample.jsonl", 50)
